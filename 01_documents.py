"""
01_documents.py — Stage 1 of the RAG pipeline: load the raw curriculum
.docx file and parse it into a structured pandas DataFrame of documents
(one row per grammar concept / literature item / reference-matrix entry).

Handles two possible curriculum.docx layouts (table-based and
paragraph-based) since the source document has changed format over time.

Run directly to sanity-check parsing against a local curriculum.docx:
    python 01_documents.py
"""

import re

import pandas as pd
import docx
from docx.oxml.ns import qn

ALL_HEADING_STYLES = ("Heading 1", "Heading 2", "Heading 3", "Heading 4")
TOP_HEADING_STYLES = ("Heading 1", "Heading 2", "Heading 3")

CONCEPT_RE = re.compile(r"^Grammar Concept:\s*(.+)$")
SCOPE_RE = re.compile(r"^Target Scope\s*&\s*Examples:\s*(.+)$")
VOCAB_RE = re.compile(r"^Key Vocabulary\s*&\s*Signals:\s*(.+)$")
ACTIVITY_RE = re.compile(r"^Suggested Activity:\s*(.+)$")
LIT_PATTERN = re.compile(r'^"([^"]+)"\s+by\s+([^:]+):\s*(.+)$')
FOCUS_GRADE_RE = re.compile(r"^Grade (\d+) Focus:$")
FOCUS_CONCEPT_RE = re.compile(r"^Key Grammar Tense\s*/\s*Concept Focus:\s*(.+)$")
FOCUS_LIT_RE = re.compile(r"^Core Assigned Literary Work\(s\):\s*(.+)$")


def _consume_plain_paragraphs(seq, i, n, stop_styles=TOP_HEADING_STYLES):
    """Collects consecutive normal paragraphs, stopping at any heading in stop_styles."""
    parts = []
    while i < n and seq[i][0] == "p" and seq[i][1] not in stop_styles:
        if seq[i][2].strip():
            parts.append(seq[i][2].strip())
        i += 1
    return parts, i


def _parse_grammar_concepts_from_paragraphs(seq, i, n):
    """Parses repeating 'Grammar Concept: / Target Scope & Examples: / Key
    Vocabulary & Signals: / Suggested Activity:' paragraph groups until the
    next heading. Returns (list_of_concept_dicts, new_index)."""
    concepts = []
    current = None
    while i < n and seq[i][0] == "p" and seq[i][1] not in ALL_HEADING_STYLES:
        text = seq[i][2].strip()
        i += 1
        if not text:
            continue
        cm = CONCEPT_RE.match(text)
        if cm:
            if current:
                concepts.append(current)
            current = {"concept": cm.group(1).strip(), "scope": "", "vocab": "", "activity": ""}
            continue
        if current is None:
            continue
        sm = SCOPE_RE.match(text)
        if sm:
            current["scope"] = sm.group(1).strip()
            continue
        vm = VOCAB_RE.match(text)
        if vm:
            current["vocab"] = vm.group(1).strip()
            continue
        am = ACTIVITY_RE.match(text)
        if am:
            current["activity"] = am.group(1).strip()
    if current:
        concepts.append(current)
    return concepts, i


def _parse_literature_from_paragraphs(seq, i, n):
    """Parses repeating '"Title" by Author: usage' lines until a non-matching
    line or a heading. Returns (list_of_(title, author, usage), new_index)."""
    books = []
    while i < n and seq[i][0] == "p" and seq[i][1] not in ALL_HEADING_STYLES:
        text = seq[i][2].strip()
        if not text:
            i += 1
            continue
        lm = LIT_PATTERN.match(text)
        if not lm:
            break
        books.append((lm.group(1), lm.group(2).strip(), lm.group(3).strip()))
        i += 1
    return books, i


def parse_curriculum_docx(path):
    d = docx.Document(path)
    body = d.element.body
    para_map = {p._p: p for p in d.paragraphs}
    table_map = {t._tbl: t for t in d.tables}

    seq = []
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            p = para_map.get(child)
            if p is not None:
                seq.append(("p", p.style.name if p.style else "normal", p.text))
        elif child.tag == qn("w:tbl"):
            t = table_map.get(child)
            if t is not None:
                rows = [[c.text.strip() for c in row.cells] for row in t.rows]
                seq.append(("tbl", None, rows))

    documents = []
    doc_id = 0
    i, n = 0, len(seq)

    while i < n:
        kind, style, content = seq[i]

        if kind == "p" and style == "Heading 2" and "Executive Summary" in content:
            i += 1
            parts, i = _consume_plain_paragraphs(seq, i, n)
            documents.append({
                "document_id": doc_id, "grade": None, "doc_type": "overview",
                "section": "Executive Summary", "grammar_topic": None, "literature": None,
                "title": "Curriculum Executive Summary",
                "text": " ".join(parts),
            })
            doc_id += 1
            continue

        if kind == "p" and style == "Heading 3":
            m = re.match(r"Grade (\d+):\s*(.*)", content)
            grade_num = int(m.group(1)) if m else None
            fallback_title = content
            i += 1

            # Overview text: plain paragraphs immediately after the grade
            # heading, stopping at ANY heading level (including Heading 4
            # subsections like "Curricular Breakdown:").
            overview_parts, i = _consume_plain_paragraphs(seq, i, n, stop_styles=ALL_HEADING_STYLES)
            documents.append({
                "document_id": doc_id, "grade": grade_num, "doc_type": "grade_overview",
                "section": "Grade Overview", "grammar_topic": None, "literature": None,
                "title": f"Grade {grade_num} Overview",
                "text": " ".join(overview_parts) if overview_parts else fallback_title,
            })
            doc_id += 1

            # --- Grammar concepts: table format OR Heading-4 paragraph format ---
            if i < n and seq[i][0] == "tbl":
                rows = seq[i][2]
                for row in rows[1:]:
                    if len(row) < 4 or not row[0].strip():
                        continue
                    concept, scope, vocab, activity = row[0], row[1], row[2], row[3]
                    documents.append({
                        "document_id": doc_id, "grade": grade_num, "doc_type": "grammar_concept",
                        "section": "Grammar Concept", "grammar_topic": concept.strip(), "literature": None,
                        "title": f"Grade {grade_num} - {concept.strip()}",
                        "text": (
                            f"Grammar concept: {concept.strip()}. "
                            f"Target scope & examples: {scope.strip()} "
                            f"Key vocabulary & signals: {vocab.strip()} "
                            f"Suggested activity: {activity.strip()}"
                        ),
                    })
                    doc_id += 1
                i += 1
            elif i < n and seq[i][0] == "p" and seq[i][1] == "Heading 4" and "Curricular Breakdown" in seq[i][2]:
                i += 1
                concepts, i = _parse_grammar_concepts_from_paragraphs(seq, i, n)
                for c in concepts:
                    documents.append({
                        "document_id": doc_id, "grade": grade_num, "doc_type": "grammar_concept",
                        "section": "Grammar Concept", "grammar_topic": c["concept"], "literature": None,
                        "title": f"Grade {grade_num} - {c['concept']}",
                        "text": (
                            f"Grammar concept: {c['concept']}. "
                            f"Target scope & examples: {c['scope']} "
                            f"Key vocabulary & signals: {c['vocab']} "
                            f"Suggested activity: {c['activity']}"
                        ),
                    })
                    doc_id += 1

            # --- Assigned literature: legacy inline marker OR Heading-4 marker ---
            if i < n and seq[i][0] == "p" and seq[i][1] == "Heading 4" and "Assigned Stories" in seq[i][2]:
                i += 1
            elif i < n and seq[i][0] == "p" and seq[i][1] not in ALL_HEADING_STYLES and "Assigned Stories" in seq[i][2]:
                i += 1

            books, i = _parse_literature_from_paragraphs(seq, i, n)
            for book_title, author, usage in books:
                documents.append({
                    "document_id": doc_id, "grade": grade_num, "doc_type": "literature",
                    "section": "Assigned Literature", "grammar_topic": None, "literature": book_title,
                    "title": f"Grade {grade_num} - {book_title}",
                    "text": f'Assigned story: "{book_title}" by {author}. {usage}',
                })
                doc_id += 1
            continue

        if kind == "p" and style == "Heading 2" and "Quick Reference Matrix" in content:
            i += 1
            intro_parts = []
            while i < n and seq[i][0] == "p" and seq[i][1] not in ALL_HEADING_STYLES:
                text = seq[i][2].strip()
                if FOCUS_GRADE_RE.match(text):
                    break  # matrix rows start here — stop treating text as intro
                if text:
                    intro_parts.append(text)
                i += 1

            lines = []
            if i < n and seq[i][0] == "tbl":
                rows = seq[i][2]
                for row in rows[1:]:
                    if len(row) < 3 or not row[0].strip():
                        continue
                    lines.append(f"{row[0].strip()}: {row[1].strip()} — literature: {row[2].strip()}")
                i += 1
            else:
                # Paragraph-triple format: "Grade N Focus:" / "Key Grammar
                # Tense / Concept Focus: ..." / "Core Assigned Literary Work(s): ..."
                pending_grade, pending_concept = None, None
                while i < n and seq[i][0] == "p" and seq[i][1] not in TOP_HEADING_STYLES:
                    text = seq[i][2].strip()
                    i += 1
                    if not text:
                        continue
                    gm = FOCUS_GRADE_RE.match(text)
                    if gm:
                        pending_grade, pending_concept = gm.group(1), None
                        continue
                    cm = FOCUS_CONCEPT_RE.match(text)
                    if cm and pending_grade is not None:
                        pending_concept = cm.group(1).strip()
                        continue
                    lm = FOCUS_LIT_RE.match(text)
                    if lm and pending_grade is not None:
                        lines.append(f"Grade {pending_grade}: {pending_concept or ''} — literature: {lm.group(1).strip()}")
                        pending_grade, pending_concept = None, None

            documents.append({
                "document_id": doc_id, "grade": None, "doc_type": "reference_matrix",
                "section": "Reference Matrix", "grammar_topic": None, "literature": None,
                "title": "Quick Reference Matrix",
                "text": " ".join(intro_parts) + " " + " | ".join(lines),
            })
            doc_id += 1
            continue

        i += 1

    df = pd.DataFrame(documents)
    if df.empty:
        raise ValueError(
            "No documents were extracted. The .docx structure may not match the "
            "expected headings (Heading 2 'Executive Summary' / Heading 3 'Grade N' / "
            "Heading 2 'Quick Reference Matrix')."
        )
    return df


if __name__ == "__main__":
    df = parse_curriculum_docx("curriculum.docx")
    print(f"Parsed {len(df)} document rows.")
    print(df.head())
