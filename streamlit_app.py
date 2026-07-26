"""
streamlit_app.py — the deployed Streamlit UI. Orchestrates the RAG pipeline
by importing each numbered stage module (01_documents.py ... 07_prompting.py)
and wires them together; also holds the app-specific bits that aren't part
of the core RAG pipeline itself (teacher login/history, coverage tracker,
worksheet exports, and the theme).

Pipeline stage modules are imported via importlib because their filenames
start with digits, which the plain `import` statement doesn't allow.
"""

import importlib
import io
import json
import os
import re
from datetime import datetime
from pathlib import Path

import docx
import streamlit as st
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

documents_stage = importlib.import_module("01_documents")
preprocessing_stage = importlib.import_module("02_preprocessing")
chunking_stage = importlib.import_module("03_chunking")
vector_representation_stage = importlib.import_module("04_vector_representation")
chroma_store_stage = importlib.import_module("05_create_chroma_store")
retrieve_context_stage = importlib.import_module("06_retrieve_context")
prompting_stage = importlib.import_module("07_prompting")

parse_curriculum_docx = documents_stage.parse_curriculum_docx
clean_text = preprocessing_stage.clean_text
build_adaptive_chunks = chunking_stage.build_adaptive_chunks
chunk_row_to_document = chunking_stage.chunk_row_to_document
get_embedding_model = vector_representation_stage.get_embedding_model
create_vector_store = chroma_store_stage.create_vector_store
retrieve_top_k = retrieve_context_stage.retrieve_top_k
build_context = retrieve_context_stage.build_context
build_grounded_prompt = prompting_stage.build_grounded_prompt
generate_answer = prompting_stage.generate_answer
build_worksheet_prompt = prompting_stage.build_worksheet_prompt

# --------------------------------------------------------------------------
# Config
# --------------------------------------------------------------------------
DOCX_PATH = "curriculum.docx"
DEFAULT_RETRIEVAL_K = 5  # was previously a "Chunks to retrieve (k)" slider in the sidebar


# --------------------------------------------------------------------------
# Cached pipeline setup (runs once per session, not on every question) —
# ties together stages 01-05: parse -> clean -> chunk -> embed -> vector store
# --------------------------------------------------------------------------
@st.cache_resource(show_spinner="Building the curriculum index…")
def build_index(docx_path):
    documents_df = parse_curriculum_docx(docx_path)
    documents_df["text"] = documents_df["text"].map(clean_text)
    chunks_df = build_adaptive_chunks(documents_df)

    embedding_model = get_embedding_model()
    all_chunk_documents = [chunk_row_to_document(row) for _, row in chunks_df.iterrows()]
    vectorstore = create_vector_store(all_chunk_documents, embedding_model)

    return chunks_df, vectorstore, embedding_model


# --------------------------------------------------------------------------
# Guided question builder — lets the user pick each part of the question
# from dropdowns (populated from the actual curriculum content) instead of
# typing free text, then composes a natural-language question from it.
# --------------------------------------------------------------------------
QUESTION_TYPES = {
    "Find which grade teaches a topic": "topic_lookup",
    "Find which grade uses a book/story": "book_lookup",
    "Find the activity for a topic in a grade": "activity",
    "List grammar topics taught in a grade": "grammar_list",
    "List stories/books taught in a grade": "book_list",
    "See everything covered in a grade": "overview",
}


def build_guided_question(chunks_df):
    """Renders cascading selectboxes and returns (question_text, grade_value).

    Framed around what a new teacher actually needs: locating which grade/
    stage a curriculum item belongs to, or finding the suggested activity
    for a topic they already know the grade of.
    """
    type_label = st.selectbox("1) What do you want to know?", list(QUESTION_TYPES.keys()))
    qtype = QUESTION_TYPES[type_label]

    if qtype == "topic_lookup":
        topics = sorted(
            chunks_df.loc[chunks_df["doc_type"] == "grammar_concept", "grammar_topic"].dropna().unique()
        )
        topic = st.selectbox("2) Grammar topic", topics)
        return f"Which grade teaches {topic}, and what stage is it introduced at?", None

    if qtype == "book_lookup":
        books = sorted(chunks_df.loc[chunks_df["doc_type"] == "literature", "literature"].dropna().unique())
        book = st.selectbox("2) Book / story", books)
        return f'Which grade uses "{book}" and at what stage is it assigned?', None

    if qtype == "activity":
        grades = sorted(chunks_df.loc[chunks_df["doc_type"] == "grammar_concept", "grade"].dropna().unique())
        grade = st.selectbox("2) Grade", [int(g) for g in grades])
        topics = sorted(
            chunks_df.loc[
                (chunks_df["doc_type"] == "grammar_concept") & (chunks_df["grade"] == grade),
                "grammar_topic",
            ].dropna().unique()
        )
        topic = st.selectbox("3) Topic", topics)
        return f"What classroom activity is suggested for teaching {topic} in grade {grade}?", grade

    if qtype == "grammar_list":
        grades = sorted(chunks_df.loc[chunks_df["doc_type"] == "grammar_concept", "grade"].dropna().unique())
        grade = st.selectbox("2) Grade", [int(g) for g in grades])
        return f"What grammar topics are taught in grade {grade}?", grade

    if qtype == "book_list":
        grades = sorted(chunks_df.loc[chunks_df["doc_type"] == "literature", "grade"].dropna().unique())
        grade = st.selectbox("2) Grade", [int(g) for g in grades])
        return f"What stories or books are assigned in grade {grade}?", grade

    # overview
    grades = sorted(chunks_df.loc[chunks_df["doc_type"] == "grade_overview", "grade"].dropna().unique())
    grade = st.selectbox("2) Grade", [int(g) for g in grades])
    return f"What topics, grammar concepts, and activities are covered in grade {grade}?", grade

TEACHER_DATA_FILE = "teacher_data.json"


def _load_all_teacher_data():
    try:
        with open(TEACHER_DATA_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {}


def _save_all_teacher_data(store):
    try:
        with open(TEACHER_DATA_FILE, "w", encoding="utf-8") as f:
            json.dump(store, f)
    except OSError:
        pass  # best-effort — some hosts have read-only/ephemeral filesystems


def load_coverage_state(teacher_name):
    store = _load_all_teacher_data()
    return store.get(teacher_name, {}).get("coverage_state", {})


def save_coverage_state(teacher_name, state):
    store = _load_all_teacher_data()
    store.setdefault(teacher_name, {})["coverage_state"] = state
    _save_all_teacher_data(store)


def load_question_history(teacher_name):
    store = _load_all_teacher_data()
    return store.get(teacher_name, {}).get("history", [])


def add_question_to_history(teacher_name, question, grade, max_items=15):
    store = _load_all_teacher_data()
    entry = store.setdefault(teacher_name, {})
    history = entry.setdefault("history", [])
    history.insert(0, {
        "question": question,
        "grade": grade,
        "asked_at": datetime.now().strftime("%b %d, %H:%M"),
    })
    entry["history"] = history[:max_items]
    _save_all_teacher_data(store)



def text_to_docx_bytes(text):
    document = docx.Document()  # new blank document (no path = create, not open)
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        else:
            document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _escape_for_reportlab(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def text_to_pdf_bytes(text):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            story.append(Spacer(1, 8))
            continue
        safe_line = _escape_for_reportlab(line)
        if safe_line.startswith("## "):
            story.append(Paragraph(safe_line[3:].strip(), styles["Heading2"]))
        elif safe_line.startswith("# "):
            story.append(Paragraph(safe_line[2:].strip(), styles["Title"]))
        else:
            story.append(Paragraph(safe_line, styles["Normal"]))
    doc.build(story)
    return buffer.getvalue()


# --------------------------------------------------------------------------
# Visual identity — "the whole app is the chalkboard": both the sidebar and
# the main pane share the same dark chalkboard green, with a slightly
# lighter green for content cards so text stays readable. Chalk-yellow is
# the one accent color, used sparingly for headings and primary actions.
# --------------------------------------------------------------------------

def inject_theme():
    st.markdown(
        """
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Bitter:wght@600;700&family=Source+Sans+3:wght@400;600&family=Caveat:wght@600&display=swap');

        :root {
            --chalkboard: #16281D;
            --chalkboard-panel: #21402C;
            --chalk: #F2EFE4;
            --chalk-dim: #C7D2C6;
            --chalk-yellow: #E8A33D;
            color-scheme: dark;
        }

        html, body, [class*="css"] { font-family: 'Source Sans 3', sans-serif; color: var(--chalk); }

        /* Main pane — same chalkboard as the sidebar. .stApp is the
           reliable top-level container across Streamlit versions. */
        .stApp {
            background-color: var(--chalkboard) !important;
            background-image: radial-gradient(rgba(255,255,255,0.03) 1px, transparent 1px);
            background-size: 6px 6px;
        }
        .stApp * {
            color: var(--chalk) !important;
        }
        .stApp, .stApp p, .stApp span, .stApp label, .stApp li,
        [data-testid="stMarkdownContainer"], [data-testid="stMarkdownContainer"] * {
            color: var(--chalk) !important;
        }
        h1, h2, h3, .stMarkdown h1, .stMarkdown h2, .stMarkdown h3 {
            font-family: 'Bitter', serif; color: var(--chalk-yellow) !important; letter-spacing: -0.01em;
        }
        .stCaption, [data-testid="stCaptionContainer"] { color: var(--chalk-dim) !important; }

        /* Content cards (expanders, text areas, etc.) sit one shade lighter
           than the board so they read as a surface on top of it */
        [data-testid="stExpander"], .stTextInput > div > div, .stSelectbox > div > div, .stTextArea textarea {
            background-color: var(--chalkboard-panel) !important; border-radius: 8px;
            border: 1px solid rgba(242,239,228,0.15) !important; color: var(--chalk) !important;
        }

        /* Sidebar — one shade lighter for gentle contrast with the main board */
        [data-testid="stSidebar"] {
            background-color: var(--chalkboard-panel);
            background-image: radial-gradient(rgba(255,255,255,0.03) 1px, transparent 1px);
            background-size: 6px 6px;
        }
        [data-testid="stSidebar"] * { color: var(--chalk) !important; }
        [data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3 {
            font-family: 'Bitter', serif; color: var(--chalk-yellow) !important;
            border-bottom: 2px dashed rgba(242,239,228,0.35); padding-bottom: 6px;
        }
        [data-testid="stSidebar"] hr { border-color: rgba(242,239,228,0.25); }

        /* Buttons */
        .stButton > button, .stDownloadButton > button {
            border-radius: 8px; border: 1px solid var(--chalk-yellow); font-weight: 600;
            background-color: transparent; color: var(--chalk-yellow) !important;
        }
        .stButton > button[kind="primary"] {
            background-color: var(--chalk-yellow); color: #16281D !important; border: none;
        }

        /* Grade badge — a chalk-drawn circle, one color per grade so grades
           are distinguishable at a glance instead of one uniform color */
        .grade-badge {
            display: inline-block; font-family: 'Caveat', cursive; font-size: 1.05rem;
            border: 2px solid currentColor; border-radius: 50%; background: transparent;
            width: 2.1em; height: 2.1em; line-height: 1.9em; text-align: center;
            transform: rotate(-6deg); margin-right: 6px;
        }

        /* Force the dark chalkboard regardless of the visitor's OS light/dark
           preference — Streamlit's own top header bar/decoration strip is
           themed separately from .stApp, so it needs its own override. */
        [data-testid="stHeader"] {
            background-color: var(--chalkboard) !important;
        }
        [data-testid="stDecoration"] {
            background-image: none !important; background-color: var(--chalkboard) !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


GRADE_BADGE_COLORS = {
    1: "#E8A33D",  # marigold
    2: "#6FA8DC",  # sky blue
    3: "#93C47D",  # leaf green
    4: "#E06666",  # coral red
    5: "#C27BA0",  # mauve
    6: "#8E7CC3",  # violet
}


GRAMMAR_FIELDS_RE = re.compile(
    r"target scope\s*&\s*examples:\s*(.*?)\s*key vocabulary\s*&\s*signals:\s*(.*?)\s*suggested activity:\s*(.*)$",
    re.IGNORECASE | re.DOTALL,
)


def grammar_concept_has_complete_content(chunk_text):
    """A grammar-concept topic only has enough material for a reliable
    worksheet/quiz answer if its scope, vocabulary, and activity fields are
    all present and non-empty in the curriculum Word file. Topics missing
    any of these (e.g. an entry with a blank scope or activity) are
    excluded from the topic dropdown rather than generating a quiz with no
    real answer behind it."""
    if not isinstance(chunk_text, str):
        return False
    match = GRAMMAR_FIELDS_RE.search(chunk_text)
    if not match:
        return False
    scope, vocab, activity = (g.strip() for g in match.groups())
    return bool(scope) and bool(vocab) and bool(activity)


def grade_badge_html(grade):
    if grade is None:
        return ""
    try:
        grade_int = int(grade)
    except (TypeError, ValueError):
        return ""
    color = GRADE_BADGE_COLORS.get(grade_int, "var(--chalk-yellow)")
    return f'<span class="grade-badge" style="color: {color} !important;">{grade_int}</span>'



# --------------------------------------------------------------------------
# Streamlit UI
# --------------------------------------------------------------------------
st.set_page_config(page_title="Curricula", page_icon="logo.png", layout="centered")
inject_theme()

logo_col, title_col = st.columns([1, 4])
with logo_col:
    st.image("logo.png", width=140)
with title_col:
    st.markdown(
        """
        <div style="line-height:1.15; margin-top: 4px;">
            <span style="font-family:'Bitter',serif; font-weight:700; font-size:2.6rem; color:var(--chalk-yellow);">Curricula</span><br>
            <span style="font-family:'Bitter',serif; font-weight:600; font-size:1.2rem; color:var(--chalk);">English Curriculum Assistant</span>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.caption("Ask a question about the Primary English Curriculum Map (Grades 1–6). Answers are grounded only in the curriculum document.")

# Token comes from Streamlit secrets / environment only — never shown or
# entered in the UI, so a using teacher never has to find or handle a token.
hf_token = os.environ.get("HF_TOKEN", st.secrets.get("HF_TOKEN", "") if hasattr(st, "secrets") else "")

# --------------------------------------------------------------------------
# Name-only "login" — no password. This is just a way to keep each
# teacher's coverage progress and question history separate on a link
# that's shared by several teachers; it is NOT real access control, since
# typing an existing teacher's name shows her data too.
# --------------------------------------------------------------------------
if "teacher_name" not in st.session_state:
    st.subheader("👋 Welcome — what's your name?")
    st.caption("This keeps your coverage progress and question history separate from other teachers using this app.")
    name_input = st.text_input("Your name")
    if st.button("Start", type="primary") and name_input.strip():
        st.session_state.teacher_name = name_input.strip()
        st.rerun()
    st.stop()

teacher_name = st.session_state.teacher_name

with st.sidebar:
    st.header("Settings")
    st.caption(f"👤 Signed in as **{teacher_name}**")
    if st.button("Switch teacher"):
        del st.session_state["teacher_name"]
        st.rerun()
    if not hf_token:
        st.warning("HF_TOKEN isn't configured — ask the app owner to add it in Streamlit secrets.", icon="⚠️")
    grade_filter = st.selectbox("Grade filter (optional)", ["Any"] + [str(g) for g in range(1, 7)])
    grade_value = None if grade_filter == "Any" else int(grade_filter)
    k = DEFAULT_RETRIEVAL_K

    st.divider()
    st.subheader("🕘 Recent questions")
    history = load_question_history(teacher_name)
    if not history:
        st.caption("No questions yet this session.")
    else:
        for i, item in enumerate(history):
            grade_tag = f" · grade {item['grade']}" if item.get("grade") else ""
            label = f"{item['question'][:45]}{'…' if len(item['question']) > 45 else ''}"
            if st.button(label, key=f"history_{i}", help=f"Asked {item['asked_at']}{grade_tag}", use_container_width=True):
                st.session_state["reuse_question"] = item["question"]
                st.rerun()

    st.divider()
    if st.button("🔄 Reset / Clear"):
        st.session_state.clear()
        st.rerun()

if not Path(DOCX_PATH).exists():
    st.error(f"'{DOCX_PATH}' not found next to app.py. Add the curriculum file to the repo.")
    st.stop()

chunks_df, vectorstore, embedding_model = build_index(DOCX_PATH)

tab_qa, tab_search, tab_worksheet, tab_coverage = st.tabs(
    ["💬 Ask a Question", "🔍 Search Curriculum", "📝 Worksheet / Quiz Generator", "✅ Coverage Tracker"]
)

# ==========================================================================
# TAB 1 — Q&A (unchanged retrieval logic)
# ==========================================================================
with tab_qa:
    reused_question = st.session_state.pop("reuse_question", None)
    if reused_question:
        st.session_state["qa_mode_radio"] = "Write a free-text question"
        st.session_state["qa_free_text"] = reused_question

    mode = st.radio(
        "Question mode", ["Pick from lists (guided)", "Write a free-text question"],
        horizontal=True, key="qa_mode_radio",
    )

    guided_grade = None
    if mode == "Pick from lists (guided)":
        question, guided_grade = build_guided_question(chunks_df)
        st.text_input("Question that will be sent", value=question, disabled=True)
    else:
        question = st.text_input(
            "Your question", placeholder="e.g. How do I teach present perfect to grade 5 students?",
            key="qa_free_text",
        )

    ask = st.button("Ask", type="primary")

    # In guided mode, the grade picked while building the question takes
    # priority over the sidebar filter (it's more specific to this question).
    effective_grade = guided_grade if mode == "Pick from lists (guided)" and guided_grade is not None else grade_value

    if ask and question.strip():
        with st.spinner("🔍 Searching the curriculum database..."):
            retrieved = retrieve_top_k(
                chunks_df, vectorstore, embedding_model, question,
                grade=effective_grade, k=k, fetch_k=max(20, k * 4),
            )
            context_text, selected_df = build_context(retrieved)
            prompt = build_grounded_prompt(question, context_text)
            answer, model_used = generate_answer(prompt, hf_token)

        add_question_to_history(teacher_name, question, effective_grade)

        st.subheader("Answer")
        st.write(answer)

        if model_used:
            st.caption(f"Generated with: {model_used}")

        if not selected_df.empty:
            with st.expander("📚 View Retrieved Sources"):
                for idx, row in selected_df.iterrows():
                    badge = grade_badge_html(row["grade"])
                    st.markdown(f"{badge}**{row['title']}**  ·  score={row['score']:.3f}", unsafe_allow_html=True)
                    st.write(row["chunk_text"])
                    if idx != selected_df.index[-1]:
                        st.divider()
        else:
            st.info("No matching chunks were retrieved for this question.")
    elif ask:
        st.warning("Type a question first.")

# ==========================================================================
# TAB — Keyword search (a plain index into the curriculum, no LLM call:
# useful when a teacher wants to browse/skim rather than ask a question)
# ==========================================================================
with tab_search:
    st.caption("Search the curriculum document directly by keyword — no AI involved, just a plain text match.")
    keyword = st.text_input("Keyword", placeholder="e.g. present perfect, Charlotte's Web, phonics…")

    if keyword.strip():
        needle = keyword.strip().lower()
        searchable_cols = [c for c in ["chunk_text", "grammar_topic", "literature", "title"] if c in chunks_df.columns]
        mask = False
        for col in searchable_cols:
            mask = mask | chunks_df[col].astype(str).str.lower().str.contains(re.escape(needle), na=False)
        if grade_value is not None:
            mask = mask & (chunks_df["grade"] == grade_value)
        results = chunks_df[mask]

        st.caption(f"{len(results)} match(es)" + (f" in grade {grade_value}" if grade_value else ""))
        for _, row in results.sort_values("grade").iterrows():
            badge = grade_badge_html(row["grade"])
            st.markdown(f"{badge}**{row['title']}**", unsafe_allow_html=True)
            st.write(row["chunk_text"])
            st.divider()
    else:
        st.info("Type a keyword above to search the whole curriculum.")

# ==========================================================================
# TAB 2 — Worksheet / Quiz generator
# ==========================================================================
with tab_worksheet:
    st.caption("Generates practice material grounded only in the selected grammar concept's scope, vocabulary, and examples.")

    grammar_df = chunks_df[chunks_df["doc_type"] == "grammar_concept"]
    grammar_df = grammar_df[grammar_df["chunk_text"].apply(grammar_concept_has_complete_content)]

    if grammar_df.empty:
        st.warning("No grammar topics with complete scope/vocabulary/activity content were found in the curriculum file.")
        st.stop()

    w_grades = sorted(grammar_df["grade"].dropna().unique())
    w_grade = st.selectbox("Grade", [int(g) for g in w_grades], key="worksheet_grade")

    w_topics = sorted(grammar_df.loc[grammar_df["grade"] == w_grade, "grammar_topic"].dropna().unique())
    w_topic = st.selectbox("Grammar topic", w_topics, key="worksheet_topic")

    w_kind_label = st.radio(
        "Type", ["Worksheet (practice only)", "Quiz (with answer key)"], horizontal=True, key="worksheet_kind"
    )
    w_kind = "quiz" if "Quiz" in w_kind_label else "worksheet"
    w_num_q = st.slider("Number of questions", min_value=3, max_value=10, value=5, key="worksheet_num_q")

    generate_worksheet = st.button("Generate", type="primary", key="worksheet_generate")

    if generate_worksheet:
        matching_rows = grammar_df[(grammar_df["grade"] == w_grade) & (grammar_df["grammar_topic"] == w_topic)]
        if matching_rows.empty:
            st.error("Couldn't find that topic's content — try a different selection.")
        else:
            topic_text = matching_rows.iloc[0]["chunk_text"]
            with st.spinner("🔍 Searching the curriculum database..."):
                prompt = build_worksheet_prompt(topic_text, w_topic, w_grade, w_kind, w_num_q)
                content, model_used = generate_answer(prompt, hf_token)

            st.subheader(f"{'Quiz' if w_kind == 'quiz' else 'Worksheet'}: {w_topic} (Grade {w_grade})")
            st.markdown(content)
            if model_used:
                st.caption(f"Generated with: {model_used}")

            export_title = f"{'Quiz' if w_kind == 'quiz' else 'Worksheet'}: {w_topic} (Grade {w_grade})"
            export_body = re.sub(r"(?im)^(answer key.*)$", r"## \1", content)
            export_text = f"# {export_title}\n\n{export_body}"
            file_stub = f"grade{w_grade}_{w_topic.replace(' ', '_')}_{w_kind}"

            dl_col1, dl_col2, dl_col3 = st.columns(3)
            with dl_col1:
                st.download_button("📥 .txt", data=content, file_name=f"{file_stub}.txt", mime="text/plain")
            with dl_col2:
                st.download_button(
                    "📥 Word (.docx)",
                    data=text_to_docx_bytes(export_text),
                    file_name=f"{file_stub}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            with dl_col3:
                st.download_button(
                    "📥 PDF", data=text_to_pdf_bytes(export_text), file_name=f"{file_stub}.pdf", mime="application/pdf"
                )

            with st.expander("📚 View Source Content"):
                st.write(topic_text)

# ==========================================================================
# TAB 3 — Coverage tracker
# ==========================================================================
with tab_coverage:
    st.caption("Track which grammar topics and stories you've already taught. Progress is saved automatically, per teacher.")

    coverage_key = f"coverage_state_{teacher_name}"
    if coverage_key not in st.session_state:
        st.session_state[coverage_key] = load_coverage_state(teacher_name)

    trackable_df = chunks_df[chunks_df["doc_type"].isin(["grammar_concept", "literature"])]
    coverage_grades = sorted(trackable_df["grade"].dropna().unique())

    total_items = len(trackable_df)
    total_done = sum(1 for cid in trackable_df["chunk_id"] if st.session_state[coverage_key].get(cid))
    if total_items:
        st.progress(total_done / total_items, text=f"Overall progress: {total_done}/{total_items} items covered")

    state_changed = False
    for g in coverage_grades:
        grade_items = trackable_df[trackable_df["grade"] == g]
        grade_done = sum(1 for cid in grade_items["chunk_id"] if st.session_state[coverage_key].get(cid))
        with st.expander(f"Grade {int(g)}  ·  {grade_done}/{len(grade_items)} covered"):
            for _, item in grade_items.iterrows():
                cid = item["chunk_id"]
                label = f"{'📖' if item['doc_type'] == 'literature' else '📘'} {item['title']}"
                current_value = st.session_state[coverage_key].get(cid, False)
                new_value = st.checkbox(label, value=current_value, key=f"cov_{teacher_name}_{cid}")
                if new_value != current_value:
                    st.session_state[coverage_key][cid] = new_value
                    state_changed = True

    if state_changed:
        save_coverage_state(teacher_name, st.session_state[coverage_key])

    st.divider()
    if st.button("🗑️ Reset coverage progress"):
        st.session_state[coverage_key] = {}
        save_coverage_state(teacher_name, {})
        st.rerun()

    st.divider()
    st.subheader("Export Coverage Report")
    report_lines = ["# Curriculum Coverage Report", "", f"Teacher: {teacher_name}", f"Overall progress: {total_done}/{total_items} items covered", ""]
    for g in coverage_grades:
        grade_items = trackable_df[trackable_df["grade"] == g]
        grade_done = sum(1 for cid in grade_items["chunk_id"] if st.session_state[coverage_key].get(cid))
        report_lines.append(f"## Grade {int(g)}  ({grade_done}/{len(grade_items)} covered)")
        for _, item in grade_items.iterrows():
            mark = "[x]" if st.session_state[coverage_key].get(item["chunk_id"]) else "[ ]"
            report_lines.append(f"{mark} {item['title']}")
        report_lines.append("")
    report_text = "\n".join(report_lines)

    rep_col1, rep_col2 = st.columns(2)
    with rep_col1:
        st.download_button(
            "📥 Word (.docx)",
            data=text_to_docx_bytes(report_text),
            file_name="curriculum_coverage_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    with rep_col2:
        st.download_button(
            "📥 PDF",
            data=text_to_pdf_bytes(report_text),
            file_name="curriculum_coverage_report.pdf",
            mime="application/pdf",
        )
