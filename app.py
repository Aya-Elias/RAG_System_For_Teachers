"""
Streamlit app — Primary English Curriculum Map RAG Assistant

Rebuilds the explicit RAG pipeline from the "v3" notebook (the version that
correctly wraps prompts in the tokenizer/chat-completion format) as a
Streamlit UI, with ONE deliberate change for deployability:

    Generation calls the HuggingFace Inference API remotely (huggingface_hub
    InferenceClient) instead of downloading and running model weights locally.
    This keeps RAM/CPU usage low enough to run on Streamlit Community Cloud's
    free tier, where loading a 3B+ parameter model locally is not feasible.

Everything upstream of generation (parsing, cleaning, chunking, metadata
extraction, embeddings, FAISS retrieval, context building) is identical in
behavior to the notebook.
"""

import io
import json
import os
import re
from datetime import datetime
from pathlib import Path

import numpy as np
import pandas as pd
import docx
from docx.oxml.ns import qn
import streamlit as st
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
from huggingface_hub import InferenceClient

# --------------------------------------------------------------------------
# Config
# --------------------------------------------------------------------------
DOCX_PATH = "curriculum.docx"
EMBEDDING_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"

# Tried in order via the Inference API. Gated Llama models still need the
# same HF license acceptance as before -- Qwen is included as a fallback
# that needs no special access approval.
LLM_CANDIDATES = [
    "meta-llama/Llama-3.2-3B-Instruct",
    "meta-llama/Llama-3.1-8B-Instruct",
    "Qwen/Qwen2.5-7B-Instruct",
]

SYSTEM_PROMPT = (
    "You are an English Curriculum Assistant. Follow the instructions in the "
    "user message exactly and stay grounded in the provided context."
)

PROMPT_TEMPLATE = """You are an English Curriculum Assistant.

Answer ONLY using the context provided below. Never invent information that is not explicitly present in the context.
If the answer cannot be found in the context, reply exactly: "I couldn't find that information in the curriculum."
Mention the Grade whenever it is available in the context.
Keep your answer concise and educational, suitable for a teacher preparing a lesson.

Context:
{context}

Question:
{question}

Answer:"""


# --------------------------------------------------------------------------
# Section 2 — Document loading
#
# Handles TWO possible curriculum.docx layouts, since the source document
# has changed format over time:
#   (a) table-based: grammar concepts / reference matrix rows live inside
#       Word tables directly under the "Grade N" / "Quick Reference Matrix"
#       headings.
#   (b) paragraph-based: no tables at all. Grammar concepts sit under a
#       "Curricular Breakdown:" (Heading 4) as repeating groups of plain
#       paragraphs ("Grammar Concept: ...", "Target Scope & Examples: ...",
#       "Key Vocabulary & Signals: ...", "Suggested Activity: ..."),
#       literature sits under an "Assigned Stories & Literature:"
#       (Heading 4), and the reference matrix is repeating paragraph
#       triples ("Grade N Focus:", "Key Grammar Tense / Concept Focus: ...",
#       "Core Assigned Literary Work(s): ...") instead of a table.
# --------------------------------------------------------------------------
ALL_HEADING_STYLES = ("Heading 1", "Heading 2", "Heading 3", "Heading 4")
TOP_HEADING_STYLES = ("Heading 1", "Heading 2", "Heading 3")

CONCEPT_RE = re.compile(r"^Grammar Concept:\s*(.+)$")
SCOPE_RE = re.compile(r"^Target Scope\s*&\s*Examples:\s*(.+)$")
VOCAB_RE = re.compile(r"^Key Vocabulary\s*&\s*Signals:\s*(.+)$")
ACTIVITY_RE = re.compile(r"^Suggested Activity:\s*(.+)$")
LIT_PATTERN = re.compile(r'^"([^"]+)"\s+by\s+([^:]+):\s*(.+)$')
FOCUS_GRADE_RE = re.compile(r"^Grade (\d+) Focus:$")
FOCUS_CONCEPT_RE = re.compile(r"^Key Grammar Tense\s*/\s*Concept Focus:\s*(.+)$")
FOCUS_LIT_RE = re.compile(r"^Core Assigned Literary Work\(s\):\s*(.+)$")


def _consume_plain_paragraphs(seq, i, n, stop_styles=TOP_HEADING_STYLES):
    """Collects consecutive normal paragraphs, stopping at any heading in stop_styles."""
    parts = []
    while i < n and seq[i][0] == "p" and seq[i][1] not in stop_styles:
        if seq[i][2].strip():
            parts.append(seq[i][2].strip())
        i += 1
    return parts, i


def _parse_grammar_concepts_from_paragraphs(seq, i, n):
    """Parses repeating 'Grammar Concept: / Target Scope & Examples: / Key
    Vocabulary & Signals: / Suggested Activity:' paragraph groups until the
    next heading. Returns (list_of_concept_dicts, new_index)."""
    concepts = []
    current = None
    while i < n and seq[i][0] == "p" and seq[i][1] not in ALL_HEADING_STYLES:
        text = seq[i][2].strip()
        i += 1
        if not text:
            continue
        cm = CONCEPT_RE.match(text)
        if cm:
            if current:
                concepts.append(current)
            current = {"concept": cm.group(1).strip(), "scope": "", "vocab": "", "activity": ""}
            continue
        if current is None:
            continue
        sm = SCOPE_RE.match(text)
        if sm:
            current["scope"] = sm.group(1).strip()
            continue
        vm = VOCAB_RE.match(text)
        if vm:
            current["vocab"] = vm.group(1).strip()
            continue
        am = ACTIVITY_RE.match(text)
        if am:
            current["activity"] = am.group(1).strip()
    if current:
        concepts.append(current)
    return concepts, i


def _parse_literature_from_paragraphs(seq, i, n):
    """Parses repeating '"Title" by Author: usage' lines until a non-matching
    line or a heading. Returns (list_of_(title, author, usage), new_index)."""
    books = []
    while i < n and seq[i][0] == "p" and seq[i][1] not in ALL_HEADING_STYLES:
        text = seq[i][2].strip()
        if not text:
            i += 1
            continue
        lm = LIT_PATTERN.match(text)
        if not lm:
            break
        books.append((lm.group(1), lm.group(2).strip(), lm.group(3).strip()))
        i += 1
    return books, i


def parse_curriculum_docx(path):
    d = docx.Document(path)
    body = d.element.body
    para_map = {p._p: p for p in d.paragraphs}
    table_map = {t._tbl: t for t in d.tables}

    seq = []
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            p = para_map.get(child)
            if p is not None:
                seq.append(("p", p.style.name if p.style else "normal", p.text))
        elif child.tag == qn("w:tbl"):
            t = table_map.get(child)
            if t is not None:
                rows = [[c.text.strip() for c in row.cells] for row in t.rows]
                seq.append(("tbl", None, rows))

    documents = []
    doc_id = 0
    i, n = 0, len(seq)

    while i < n:
        kind, style, content = seq[i]

        if kind == "p" and style == "Heading 2" and "Executive Summary" in content:
            i += 1
            parts, i = _consume_plain_paragraphs(seq, i, n)
            documents.append({
                "document_id": doc_id, "grade": None, "doc_type": "overview",
                "section": "Executive Summary", "grammar_topic": None, "literature": None,
                "title": "Curriculum Executive Summary",
                "text": " ".join(parts),
            })
            doc_id += 1
            continue

        if kind == "p" and style == "Heading 3":
            m = re.match(r"Grade (\d+):\s*(.*)", content)
            grade_num = int(m.group(1)) if m else None
            fallback_title = content
            i += 1

            # Overview text: plain paragraphs immediately after the grade
            # heading, stopping at ANY heading level (including Heading 4
            # subsections like "Curricular Breakdown:").
            overview_parts, i = _consume_plain_paragraphs(seq, i, n, stop_styles=ALL_HEADING_STYLES)
            documents.append({
                "document_id": doc_id, "grade": grade_num, "doc_type": "grade_overview",
                "section": "Grade Overview", "grammar_topic": None, "literature": None,
                "title": f"Grade {grade_num} Overview",
                "text": " ".join(overview_parts) if overview_parts else fallback_title,
            })
            doc_id += 1

            # --- Grammar concepts: table format OR Heading-4 paragraph format ---
            if i < n and seq[i][0] == "tbl":
                rows = seq[i][2]
                for row in rows[1:]:
                    if len(row) < 4 or not row[0].strip():
                        continue
                    concept, scope, vocab, activity = row[0], row[1], row[2], row[3]
                    documents.append({
                        "document_id": doc_id, "grade": grade_num, "doc_type": "grammar_concept",
                        "section": "Grammar Concept", "grammar_topic": concept.strip(), "literature": None,
                        "title": f"Grade {grade_num} - {concept.strip()}",
                        "text": (
                            f"Grammar concept: {concept.strip()}. "
                            f"Target scope & examples: {scope.strip()} "
                            f"Key vocabulary & signals: {vocab.strip()} "
                            f"Suggested activity: {activity.strip()}"
                        ),
                    })
                    doc_id += 1
                i += 1
            elif i < n and seq[i][0] == "p" and seq[i][1] == "Heading 4" and "Curricular Breakdown" in seq[i][2]:
                i += 1
                concepts, i = _parse_grammar_concepts_from_paragraphs(seq, i, n)
                for c in concepts:
                    documents.append({
                        "document_id": doc_id, "grade": grade_num, "doc_type": "grammar_concept",
                        "section": "Grammar Concept", "grammar_topic": c["concept"], "literature": None,
                        "title": f"Grade {grade_num} - {c['concept']}",
                        "text": (
                            f"Grammar concept: {c['concept']}. "
                            f"Target scope & examples: {c['scope']} "
                            f"Key vocabulary & signals: {c['vocab']} "
                            f"Suggested activity: {c['activity']}"
                        ),
                    })
                    doc_id += 1

            # --- Assigned literature: legacy inline marker OR Heading-4 marker ---
            if i < n and seq[i][0] == "p" and seq[i][1] == "Heading 4" and "Assigned Stories" in seq[i][2]:
                i += 1
            elif i < n and seq[i][0] == "p" and seq[i][1] not in ALL_HEADING_STYLES and "Assigned Stories" in seq[i][2]:
                i += 1

            books, i = _parse_literature_from_paragraphs(seq, i, n)
            for book_title, author, usage in books:
                documents.append({
                    "document_id": doc_id, "grade": grade_num, "doc_type": "literature",
                    "section": "Assigned Literature", "grammar_topic": None, "literature": book_title,
                    "title": f"Grade {grade_num} - {book_title}",
                    "text": f'Assigned story: "{book_title}" by {author}. {usage}',
                })
                doc_id += 1
            continue

        if kind == "p" and style == "Heading 2" and "Quick Reference Matrix" in content:
            i += 1
            intro_parts = []
            while i < n and seq[i][0] == "p" and seq[i][1] not in ALL_HEADING_STYLES:
                text = seq[i][2].strip()
                if FOCUS_GRADE_RE.match(text):
                    break  # matrix rows start here — stop treating text as intro
                if text:
                    intro_parts.append(text)
                i += 1

            lines = []
            if i < n and seq[i][0] == "tbl":
                rows = seq[i][2]
                for row in rows[1:]:
                    if len(row) < 3 or not row[0].strip():
                        continue
                    lines.append(f"{row[0].strip()}: {row[1].strip()} — literature: {row[2].strip()}")
                i += 1
            else:
                # Paragraph-triple format: "Grade N Focus:" / "Key Grammar
                # Tense / Concept Focus: ..." / "Core Assigned Literary Work(s): ..."
                pending_grade, pending_concept = None, None
                while i < n and seq[i][0] == "p" and seq[i][1] not in TOP_HEADING_STYLES:
                    text = seq[i][2].strip()
                    i += 1
                    if not text:
                        continue
                    gm = FOCUS_GRADE_RE.match(text)
                    if gm:
                        pending_grade, pending_concept = gm.group(1), None
                        continue
                    cm = FOCUS_CONCEPT_RE.match(text)
                    if cm and pending_grade is not None:
                        pending_concept = cm.group(1).strip()
                        continue
                    lm = FOCUS_LIT_RE.match(text)
                    if lm and pending_grade is not None:
                        lines.append(f"Grade {pending_grade}: {pending_concept or ''} — literature: {lm.group(1).strip()}")
                        pending_grade, pending_concept = None, None

            documents.append({
                "document_id": doc_id, "grade": None, "doc_type": "reference_matrix",
                "section": "Reference Matrix", "grammar_topic": None, "literature": None,
                "title": "Quick Reference Matrix",
                "text": " ".join(intro_parts) + " " + " | ".join(lines),
            })
            doc_id += 1
            continue

        i += 1

    df = pd.DataFrame(documents)
    if df.empty:
        raise ValueError(
            "No documents were extracted. The .docx structure may not match the "
            "expected headings (Heading 2 'Executive Summary' / Heading 3 'Grade N' / "
            "Heading 2 'Quick Reference Matrix')."
        )
    return df


# --------------------------------------------------------------------------
# Section 3 — Cleaning
# --------------------------------------------------------------------------
def clean_text(text):
    if not isinstance(text, str):
        return text
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"[-_=]{3,}", " ", text)
    text = re.sub(r"\s+([,.;:])", r"\1", text)
    text = re.sub(r"\s*\n\s*", " ", text)
    return text.strip()


# --------------------------------------------------------------------------
# Section 4 — Adaptive chunking
# --------------------------------------------------------------------------
CONCEPT_TYPES = {"grade_overview", "grammar_concept", "literature"}
MAX_WORDS_PER_CHUNK = 160


def split_by_structure(text, max_words=MAX_WORDS_PER_CHUNK):
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    parts, current, count = [], [], 0
    for s in sentences:
        w = len(s.split())
        if current and count + w > max_words:
            parts.append(" ".join(current))
            current, count = [], 0
        current.append(s)
        count += w
    if current:
        parts.append(" ".join(current))
    return parts if parts else [text]


def build_adaptive_chunks(documents_frame):
    rows = []
    for _, doc in documents_frame.iterrows():
        if doc["doc_type"] in CONCEPT_TYPES:
            pieces = [doc["text"]]
        else:
            pieces = split_by_structure(doc["text"])
        for idx, piece in enumerate(pieces):
            rows.append({
                "chunk_id": f"doc{doc['document_id']}_chunk{idx}",
                "document_id": doc["document_id"],
                "grade": doc["grade"],
                "section": doc["section"],
                "doc_type": doc["doc_type"],
                "grammar_topic": doc["grammar_topic"],
                "literature": doc["literature"],
                "title": doc["title"],
                "chunk_text": piece,
                "word_count": len(piece.split()),
            })
    return pd.DataFrame(rows)


def chunk_row_to_document(row):
    return Document(
        page_content=row["chunk_text"],
        metadata={
            "chunk_id": row["chunk_id"],
            "grade": row["grade"],
            "section": row["section"],
            "grammar_topic": row["grammar_topic"],
            "literature": row["literature"],
            "title": row["title"],
        },
    )


# --------------------------------------------------------------------------
# Section 8 — Metadata filtering + MMR retrieval
# --------------------------------------------------------------------------
def filter_chunks_by_metadata(df, grade=None, section=None):
    subset = df
    if grade is not None and not (isinstance(grade, float) and np.isnan(grade)):
        grade_mask = subset["grade"] == grade
        if grade_mask.any():
            subset = subset[grade_mask]
    if section:
        section_mask = subset["section"] == section
        if section_mask.any():
            subset = subset[section_mask]
    return subset.reset_index(drop=True)


def retrieve_top_k(chunks_df, vectorstore, embedding_model, query, grade=None, section=None, k=5, fetch_k=20):
    filtered_df = filter_chunks_by_metadata(chunks_df, grade=grade, section=section)
    if filtered_df.empty:
        filtered_df = chunks_df
        local_store = vectorstore
    else:
        local_docs = [chunk_row_to_document(row) for _, row in filtered_df.iterrows()]
        local_store = FAISS.from_documents(local_docs, embedding_model)

    k_eff = min(k, local_store.index.ntotal)
    fetch_k_eff = min(max(fetch_k, k_eff), local_store.index.ntotal)

    retriever = local_store.as_retriever(
        search_type="mmr",
        search_kwargs={"k": k_eff, "fetch_k": fetch_k_eff},
    )
    retrieved_docs = retriever.invoke(query)

    scored = local_store.similarity_search_with_relevance_scores(query, k=k_eff)
    score_by_chunk_id = {doc.metadata["chunk_id"]: score for doc, score in scored}

    rows = []
    for doc in retrieved_docs:
        rows.append({
            "chunk_id": doc.metadata["chunk_id"],
            "grade": doc.metadata["grade"],
            "section": doc.metadata["section"],
            "grammar_topic": doc.metadata["grammar_topic"],
            "literature": doc.metadata["literature"],
            "title": doc.metadata["title"],
            "chunk_text": doc.page_content,
            "score": score_by_chunk_id.get(doc.metadata["chunk_id"], 0.0),
        })
    return pd.DataFrame(rows)


# --------------------------------------------------------------------------
# Section 9 — Context builder
# --------------------------------------------------------------------------
def build_context(retrieved_df, word_budget=220, max_chunks=4):
    if retrieved_df.empty:
        return "", pd.DataFrame()

    ranked = retrieved_df.sort_values(by="score", ascending=False).reset_index(drop=True)

    selected_rows, seen_texts, used_words = [], set(), 0
    for _, row in ranked.iterrows():
        normalized_text = re.sub(r"\s+", " ", row["chunk_text"]).strip().lower()
        if normalized_text in seen_texts:
            continue
        chunk_words = len(row["chunk_text"].split())
        if selected_rows and used_words + chunk_words > word_budget:
            continue
        selected_rows.append(row.to_dict())
        seen_texts.add(normalized_text)
        used_words += chunk_words
        if len(selected_rows) >= max_chunks:
            break

    selected_df = pd.DataFrame(selected_rows)
    context_blocks = []
    for idx, row in enumerate(selected_rows, start=1):
        header = f"[Source {idx}] {row['title']} (Grade={row['grade']}, Section={row['section']})"
        context_blocks.append(f"{header}\n{row['chunk_text']}")
    return "\n\n".join(context_blocks), selected_df


def build_grounded_prompt(query, context_text):
    if not context_text:
        context_text = "(No relevant context was retrieved.)"
    return PROMPT_TEMPLATE.format(context=context_text, question=query)


def extract_final_answer(raw_text):
    if not raw_text:
        return ""
    cleaned = re.sub(r"<think>.*?</think>", "", raw_text, flags=re.DOTALL | re.IGNORECASE)
    cleaned = re.sub(r"\n\s*\n+", "\n\n", cleaned).strip()
    return cleaned


# --------------------------------------------------------------------------
# Generation via HF Inference API (remote — no local weights)
# --------------------------------------------------------------------------
def generate_answer(prompt, hf_token):
    """Try each candidate model via the Inference API chat-completion
    endpoint until one succeeds. Returns (answer_text, model_used_or_None)."""
    if not hf_token:
        return (
            "[No LLM response — no HuggingFace token was configured. "
            "Ask the app owner to add HF_TOKEN in Streamlit secrets.]",
            None,
        )

    last_error = None
    for candidate in LLM_CANDIDATES:
        try:
            client = InferenceClient(model=candidate, token=hf_token)
            completion = client.chat_completion(
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": prompt},
                ],
                max_tokens=250,
                temperature=0.0,
            )
            answer = completion.choices[0].message.content
            return extract_final_answer(answer), candidate
        except Exception as exc:
            last_error = exc
            continue

    return (
        f"[No LLM response — all candidate models failed. "
        f"Last error: {last_error}]",
        None,
    )



# --------------------------------------------------------------------------
# Cached pipeline setup (runs once per session, not on every question)
# --------------------------------------------------------------------------
@st.cache_resource(show_spinner="Building the curriculum index…")
def build_index(docx_path):
    documents_df = parse_curriculum_docx(docx_path)
    documents_df["text"] = documents_df["text"].map(clean_text)
    chunks_df = build_adaptive_chunks(documents_df)

    embedding_model = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
    all_chunk_documents = [chunk_row_to_document(row) for _, row in chunks_df.iterrows()]
    vectorstore = FAISS.from_documents(all_chunk_documents, embedding_model)

    return chunks_df, vectorstore, embedding_model


# --------------------------------------------------------------------------
# Guided question builder — lets the user pick each part of the question
# from dropdowns (populated from the actual curriculum content) instead of
# typing free text, then composes a natural-language question from it.
# --------------------------------------------------------------------------
QUESTION_TYPES = {
    "Find which grade teaches a topic": "topic_lookup",
    "Find which grade uses a book/story": "book_lookup",
    "Find the activity for a topic in a grade": "activity",
    "List grammar topics taught in a grade": "grammar_list",
    "List stories/books taught in a grade": "book_list",
    "See everything covered in a grade": "overview",
}


def build_guided_question(chunks_df):
    """Renders cascading selectboxes and returns (question_text, grade_value).

    Framed around what a new teacher actually needs: locating which grade/
    stage a curriculum item belongs to, or finding the suggested activity
    for a topic they already know the grade of.
    """
    type_label = st.selectbox("1) What do you want to know?", list(QUESTION_TYPES.keys()))
    qtype = QUESTION_TYPES[type_label]

    if qtype == "topic_lookup":
        topics = sorted(
            chunks_df.loc[chunks_df["doc_type"] == "grammar_concept", "grammar_topic"].dropna().unique()
        )
        topic = st.selectbox("2) Grammar topic", topics)
        return f"Which grade teaches {topic}, and what stage is it introduced at?", None

    if qtype == "book_lookup":
        books = sorted(chunks_df.loc[chunks_df["doc_type"] == "literature", "literature"].dropna().unique())
        book = st.selectbox("2) Book / story", books)
        return f'Which grade uses "{book}" and at what stage is it assigned?', None

    if qtype == "activity":
        grades = sorted(chunks_df.loc[chunks_df["doc_type"] == "grammar_concept", "grade"].dropna().unique())
        grade = st.selectbox("2) Grade", [int(g) for g in grades])
        topics = sorted(
            chunks_df.loc[
                (chunks_df["doc_type"] == "grammar_concept") & (chunks_df["grade"] == grade),
                "grammar_topic",
            ].dropna().unique()
        )
        topic = st.selectbox("3) Topic", topics)
        return f"What classroom activity is suggested for teaching {topic} in grade {grade}?", grade

    if qtype == "grammar_list":
        grades = sorted(chunks_df.loc[chunks_df["doc_type"] == "grammar_concept", "grade"].dropna().unique())
        grade = st.selectbox("2) Grade", [int(g) for g in grades])
        return f"What grammar topics are taught in grade {grade}?", grade

    if qtype == "book_list":
        grades = sorted(chunks_df.loc[chunks_df["doc_type"] == "literature", "grade"].dropna().unique())
        grade = st.selectbox("2) Grade", [int(g) for g in grades])
        return f"What stories or books are assigned in grade {grade}?", grade

    # overview
    grades = sorted(chunks_df.loc[chunks_df["doc_type"] == "grade_overview", "grade"].dropna().unique())
    grade = st.selectbox("2) Grade", [int(g) for g in grades])
    return f"What topics, grammar concepts, and activities are covered in grade {grade}?", grade


# --------------------------------------------------------------------------
# Worksheet / Quiz generator
# --------------------------------------------------------------------------
WORKSHEET_PROMPT_TEMPLATE = """You are an English Curriculum Assistant creating classroom materials for a teacher.

Using ONLY the grammar concept information below, create a {kind} for grade {grade} students on "{topic}".

Grammar concept information:
{context}

Requirements:
- Write exactly {num_questions} questions, numbered 1 to {num_questions}.
- Base every question strictly on the scope, vocabulary, and examples given above. Do not introduce grammar points that are not mentioned in the information above.
- Use simple, age-appropriate language for grade {grade} students.
- Vary the question types where appropriate (fill-in-the-blank, multiple choice, short answer).
{answer_key_instruction}

Output format:
- A short title line
- The numbered questions
{answer_key_label}"""


def build_worksheet_prompt(topic_text, topic, grade, kind, num_questions):
    is_quiz = kind == "quiz"
    answer_key_instruction = (
        "- After the questions, include an answer key."
        if is_quiz
        else "- Do NOT include an answer key — this is for independent student practice."
    )
    answer_key_label = "- An 'Answer Key' section at the end" if is_quiz else ""
    return WORKSHEET_PROMPT_TEMPLATE.format(
        kind="quiz (with an answer key)" if is_quiz else "practice worksheet (no answer key)",
        grade=grade,
        topic=topic,
        context=topic_text,
        num_questions=num_questions,
        answer_key_instruction=answer_key_instruction,
        answer_key_label=answer_key_label,
    )


# --------------------------------------------------------------------------
# --------------------------------------------------------------------------
# Per-teacher persistence — a simple JSON file keyed by teacher name, so
# coverage progress and question history don't leak between teachers
# sharing the same deployed link. This is name-only "login" with no
# password, so it is NOT real access control — anyone who types an
# existing teacher's name sees her data. It only prevents different
# teachers' normal usage from overwriting each other by accident.
# --------------------------------------------------------------------------
TEACHER_DATA_FILE = "teacher_data.json"


def _load_all_teacher_data():
    try:
        with open(TEACHER_DATA_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {}


def _save_all_teacher_data(store):
    try:
        with open(TEACHER_DATA_FILE, "w", encoding="utf-8") as f:
            json.dump(store, f)
    except OSError:
        pass  # best-effort — some hosts have read-only/ephemeral filesystems


def load_coverage_state(teacher_name):
    store = _load_all_teacher_data()
    return store.get(teacher_name, {}).get("coverage_state", {})


def save_coverage_state(teacher_name, state):
    store = _load_all_teacher_data()
    store.setdefault(teacher_name, {})["coverage_state"] = state
    _save_all_teacher_data(store)


def load_question_history(teacher_name):
    store = _load_all_teacher_data()
    return store.get(teacher_name, {}).get("history", [])


def add_question_to_history(teacher_name, question, grade, max_items=15):
    store = _load_all_teacher_data()
    entry = store.setdefault(teacher_name, {})
    history = entry.setdefault("history", [])
    history.insert(0, {
        "question": question,
        "grade": grade,
        "asked_at": datetime.now().strftime("%b %d, %H:%M"),
    })
    entry["history"] = history[:max_items]
    _save_all_teacher_data(store)


# --------------------------------------------------------------------------
# Document export (Word / PDF) — shared by the worksheet generator and the
# coverage tracker report. Input text uses a simple markdown-like
# convention: a line starting with "# " becomes a title, "## " becomes a
# section heading, everything else is a normal paragraph.
# --------------------------------------------------------------------------
def text_to_docx_bytes(text):
    document = docx.Document()  # new blank document (no path = create, not open)
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        else:
            document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _escape_for_reportlab(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def text_to_pdf_bytes(text):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            story.append(Spacer(1, 8))
            continue
        safe_line = _escape_for_reportlab(line)
        if safe_line.startswith("## "):
            story.append(Paragraph(safe_line[3:].strip(), styles["Heading2"]))
        elif safe_line.startswith("# "):
            story.append(Paragraph(safe_line[2:].strip(), styles["Title"]))
        else:
            story.append(Paragraph(safe_line, styles["Normal"]))
    doc.build(story)
    return buffer.getvalue()


# --------------------------------------------------------------------------
# Visual identity — "the whole app is the chalkboard": both the sidebar and
# the main pane share the same dark chalkboard green, with a slightly
# lighter green for content cards so text stays readable. Chalk-yellow is
# the one accent color, used sparingly for headings and primary actions.
# --------------------------------------------------------------------------
def inject_theme():
    st.markdown(
        """
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Bitter:wght@600;700&family=Source+Sans+3:wght@400;600&family=Caveat:wght@600&family=Kalam:wght@700&display=swap');

        :root {
            --chalkboard: #16281D;
            --chalkboard-panel: #21402C;
            --chalk: #F2EFE4;
            --chalk-dim: #C7D2C6;
            --chalk-yellow: #E8A33D;
        }

        html, body, [class*="css"] { font-family: 'Source Sans 3', sans-serif; color: var(--chalk); }

        h1, h2, h3, .stMarkdown h1, .stMarkdown h2, .stMarkdown h3 {
            font-family: 'Bitter', serif; color: var(--chalk-yellow); letter-spacing: -0.01em;
        }

        /* Main pane — same chalkboard as the sidebar. .stApp is the
           reliable top-level container across Streamlit versions. */
        .stApp {
            background-color: var(--chalkboard) !important;
            background-image: radial-gradient(rgba(255,255,255,0.03) 1px, transparent 1px);
            background-size: 6px 6px;
        }
        .stApp, .stApp p, .stApp span, .stApp label, .stApp li,
        [data-testid="stMarkdownContainer"], [data-testid="stMarkdownContainer"] * {
            color: var(--chalk) !important;
        }
        .stCaption, [data-testid="stCaptionContainer"] { color: var(--chalk-dim) !important; }

        /* Content cards (expanders, text areas, etc.) sit one shade lighter
           than the board so they read as a surface on top of it */
        [data-testid="stExpander"], .stTextInput > div > div, .stSelectbox > div > div, .stTextArea textarea {
            background-color: var(--chalkboard-panel) !important; border-radius: 8px;
            border: 1px solid rgba(242,239,228,0.15) !important; color: var(--chalk) !important;
        }

        /* Sidebar — one shade lighter for gentle contrast with the main board */
        [data-testid="stSidebar"] {
            background-color: var(--chalkboard-panel);
            background-image: radial-gradient(rgba(255,255,255,0.03) 1px, transparent 1px);
            background-size: 6px 6px;
        }
        [data-testid="stSidebar"] * { color: var(--chalk) !important; }
        [data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3 {
            font-family: 'Bitter', serif; color: var(--chalk-yellow) !important;
            border-bottom: 2px dashed rgba(242,239,228,0.35); padding-bottom: 6px;
        }
        [data-testid="stSidebar"] hr { border-color: rgba(242,239,228,0.25); }

        /* Buttons */
        .stButton > button, .stDownloadButton > button {
            border-radius: 8px; border: 1px solid var(--chalk-yellow); font-weight: 600;
            background-color: transparent; color: var(--chalk-yellow) !important;
        }
        .stButton > button[kind="primary"] {
            background-color: var(--chalk-yellow); color: #16281D !important; border: none;
        }

        /* Grade badge — a chalk-drawn circle, one color per grade so grades
           are distinguishable at a glance instead of one uniform color */
        .grade-badge {
            display: inline-block; font-family: 'Caveat', cursive; font-size: 1.05rem;
            border: 2px solid currentColor; border-radius: 50%; background: transparent;
            width: 2.1em; height: 2.1em; line-height: 1.9em; text-align: center;
            transform: rotate(-6deg); margin-right: 6px;
        }

        /* Force the dark chalkboard regardless of the visitor's OS light/dark
           preference — Streamlit's own top header bar/decoration strip is
           themed separately from .stApp, so it needs its own override. */
        /* Hero title/caption — styled to look hand-written in chalk */
        .chalk-hero h1 {
            font-family: 'Kalam', cursive; font-weight: 700; color: var(--chalk);
            font-size: 2.4rem; letter-spacing: 0.5px; margin-bottom: 0.2em;
            text-shadow:
                0 0 1px rgba(242,239,228,0.7),
                0 0 6px rgba(242,239,228,0.35),
                1px 1px 0 rgba(242,239,228,0.12);
        }
        .chalk-hero p {
            font-family: 'Kalam', cursive; color: var(--chalk-dim) !important;
            font-size: 1.05rem; text-shadow: 0 0 4px rgba(242,239,228,0.2);
        }

        [data-testid="stHeader"] {
            background-color: var(--chalkboard) !important;
        }
        [data-testid="stDecoration"] {
            background-image: none !important; background-color: var(--chalkboard) !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


GRADE_BADGE_COLORS = {
    1: "#E8A33D",  # marigold
    2: "#6FA8DC",  # sky blue
    3: "#93C47D",  # leaf green
    4: "#E06666",  # coral red
    5: "#C27BA0",  # mauve
    6: "#8E7CC3",  # violet
}


def grade_badge_html(grade):
    if grade is None:
        return ""
    try:
        grade_int = int(grade)
    except (TypeError, ValueError):
        return ""
    color = GRADE_BADGE_COLORS.get(grade_int, "var(--chalk-yellow)")
    return f'<span class="grade-badge" style="color: {color};">{grade_int}</span>'


# --------------------------------------------------------------------------
# Streamlit UI
# --------------------------------------------------------------------------
st.set_page_config(page_title="Curriculum RAG Assistant", page_icon="📚", layout="centered")
inject_theme()
st.markdown(
    """
    <div class="chalk-hero">
        <h1>📚 English Curriculum Assistant</h1>
        <p>Ask a question about the Primary English Curriculum Map (Grades 1–6). Answers are grounded only in the curriculum document.</p>
    </div>
    """,
    unsafe_allow_html=True,
)

# Token comes from Streamlit secrets / environment only — never shown or
# entered in the UI, so a using teacher never has to find or handle a token.
hf_token = os.environ.get("HF_TOKEN", st.secrets.get("HF_TOKEN", "") if hasattr(st, "secrets") else "")

# --------------------------------------------------------------------------
# Name-only "login" — no password. This is just a way to keep each
# teacher's coverage progress and question history separate on a link
# that's shared by several teachers; it is NOT real access control, since
# typing an existing teacher's name shows her data too.
# --------------------------------------------------------------------------
if "teacher_name" not in st.session_state:
    st.subheader("👋 Welcome — what's your name?")
    st.caption("This keeps your coverage progress and question history separate from other teachers using this app.")
    name_input = st.text_input("Your name")
    if st.button("Start", type="primary") and name_input.strip():
        st.session_state.teacher_name = name_input.strip()
        st.rerun()
    st.stop()

teacher_name = st.session_state.teacher_name

with st.sidebar:
    st.header("Settings")
    st.caption(f"👤 Signed in as **{teacher_name}**")
    if st.button("Switch teacher"):
        del st.session_state["teacher_name"]
        st.rerun()
    if not hf_token:
        st.warning("HF_TOKEN isn't configured — ask the app owner to add it in Streamlit secrets.", icon="⚠️")
    grade_filter = st.selectbox("Grade filter (optional)", ["Any"] + [str(g) for g in range(1, 7)])
    grade_value = None if grade_filter == "Any" else int(grade_filter)
    k = st.slider("Chunks to retrieve (k)", min_value=2, max_value=8, value=5)

    st.divider()
    st.subheader("🕘 Recent questions")
    history = load_question_history(teacher_name)
    if not history:
        st.caption("No questions yet this session.")
    else:
        for i, item in enumerate(history):
            grade_tag = f" · grade {item['grade']}" if item.get("grade") else ""
            label = f"{item['question'][:45]}{'…' if len(item['question']) > 45 else ''}"
            if st.button(label, key=f"history_{i}", help=f"Asked {item['asked_at']}{grade_tag}", use_container_width=True):
                st.session_state["reuse_question"] = item["question"]
                st.rerun()

    st.divider()
    if st.button("🔄 Reset / Clear"):
        st.session_state.clear()
        st.rerun()

if not Path(DOCX_PATH).exists():
    st.error(f"'{DOCX_PATH}' not found next to app.py. Add the curriculum file to the repo.")
    st.stop()

chunks_df, vectorstore, embedding_model = build_index(DOCX_PATH)

tab_qa, tab_search, tab_worksheet, tab_coverage = st.tabs(
    ["💬 Ask a Question", "🔍 Search Curriculum", "📝 Worksheet / Quiz Generator", "✅ Coverage Tracker"]
)

# ==========================================================================
# TAB 1 — Q&A (unchanged retrieval logic)
# ==========================================================================
with tab_qa:
    reused_question = st.session_state.pop("reuse_question", None)
    if reused_question:
        st.session_state["qa_mode_radio"] = "Write a free-text question"
        st.session_state["qa_free_text"] = reused_question

    mode = st.radio(
        "Question mode", ["Pick from lists (guided)", "Write a free-text question"],
        horizontal=True, key="qa_mode_radio",
    )

    guided_grade = None
    if mode == "Pick from lists (guided)":
        question, guided_grade = build_guided_question(chunks_df)
        st.text_input("Question that will be sent", value=question, disabled=True)
    else:
        question = st.text_input(
            "Your question", placeholder="e.g. How do I teach present perfect to grade 5 students?",
            key="qa_free_text",
        )

    ask = st.button("Ask", type="primary")

    # In guided mode, the grade picked while building the question takes
    # priority over the sidebar filter (it's more specific to this question).
    effective_grade = guided_grade if mode == "Pick from lists (guided)" and guided_grade is not None else grade_value

    if ask and question.strip():
        with st.spinner("🔍 Searching the curriculum database..."):
            retrieved = retrieve_top_k(
                chunks_df, vectorstore, embedding_model, question,
                grade=effective_grade, k=k, fetch_k=max(20, k * 4),
            )
            context_text, selected_df = build_context(retrieved)
            prompt = build_grounded_prompt(question, context_text)
            answer, model_used = generate_answer(prompt, hf_token)

        add_question_to_history(teacher_name, question, effective_grade)

        st.subheader("Answer")
        st.write(answer)

        if model_used:
            st.caption(f"Generated with: {model_used}")

        if not selected_df.empty:
            with st.expander("📚 View Retrieved Sources"):
                for idx, row in selected_df.iterrows():
                    badge = grade_badge_html(row["grade"])
                    st.markdown(f"{badge}**{row['title']}**  ·  score={row['score']:.3f}", unsafe_allow_html=True)
                    st.write(row["chunk_text"])
                    if idx != selected_df.index[-1]:
                        st.divider()
        else:
            st.info("No matching chunks were retrieved for this question.")
    elif ask:
        st.warning("Type a question first.")

# ==========================================================================
# TAB — Keyword search (a plain index into the curriculum, no LLM call:
# useful when a teacher wants to browse/skim rather than ask a question)
# ==========================================================================
with tab_search:
    st.caption("Search the curriculum document directly by keyword — no AI involved, just a plain text match.")
    keyword = st.text_input("Keyword", placeholder="e.g. present perfect, Charlotte's Web, phonics…")

    if keyword.strip():
        needle = keyword.strip().lower()
        searchable_cols = [c for c in ["chunk_text", "grammar_topic", "literature", "title"] if c in chunks_df.columns]
        mask = False
        for col in searchable_cols:
            mask = mask | chunks_df[col].astype(str).str.lower().str.contains(re.escape(needle), na=False)
        if grade_value is not None:
            mask = mask & (chunks_df["grade"] == grade_value)
        results = chunks_df[mask]

        st.caption(f"{len(results)} match(es)" + (f" in grade {grade_value}" if grade_value else ""))
        for _, row in results.sort_values("grade").iterrows():
            badge = grade_badge_html(row["grade"])
            st.markdown(f"{badge}**{row['title']}**", unsafe_allow_html=True)
            st.write(row["chunk_text"])
            st.divider()
    else:
        st.info("Type a keyword above to search the whole curriculum.")

# ==========================================================================
# TAB 2 — Worksheet / Quiz generator
# ==========================================================================
with tab_worksheet:
    st.caption("Generates practice material grounded only in the selected grammar concept's scope, vocabulary, and examples.")

    grammar_df = chunks_df[chunks_df["doc_type"] == "grammar_concept"]
    w_grades = sorted(grammar_df["grade"].dropna().unique())
    w_grade = st.selectbox("Grade", [int(g) for g in w_grades], key="worksheet_grade")

    w_topics = sorted(grammar_df.loc[grammar_df["grade"] == w_grade, "grammar_topic"].dropna().unique())
    w_topic = st.selectbox("Grammar topic", w_topics, key="worksheet_topic")

    w_kind_label = st.radio(
        "Type", ["Worksheet (practice only)", "Quiz (with answer key)"], horizontal=True, key="worksheet_kind"
    )
    w_kind = "quiz" if "Quiz" in w_kind_label else "worksheet"
    w_num_q = st.slider("Number of questions", min_value=3, max_value=10, value=5, key="worksheet_num_q")

    generate_worksheet = st.button("Generate", type="primary", key="worksheet_generate")

    if generate_worksheet:
        matching_rows = grammar_df[(grammar_df["grade"] == w_grade) & (grammar_df["grammar_topic"] == w_topic)]
        if matching_rows.empty:
            st.error("Couldn't find that topic's content — try a different selection.")
        else:
            topic_text = matching_rows.iloc[0]["chunk_text"]
            with st.spinner("🔍 Searching the curriculum database..."):
                prompt = build_worksheet_prompt(topic_text, w_topic, w_grade, w_kind, w_num_q)
                content, model_used = generate_answer(prompt, hf_token)

            st.subheader(f"{'Quiz' if w_kind == 'quiz' else 'Worksheet'}: {w_topic} (Grade {w_grade})")
            st.markdown(content)
            if model_used:
                st.caption(f"Generated with: {model_used}")

            export_title = f"{'Quiz' if w_kind == 'quiz' else 'Worksheet'}: {w_topic} (Grade {w_grade})"
            export_body = re.sub(r"(?im)^(answer key.*)$", r"## \1", content)
            export_text = f"# {export_title}\n\n{export_body}"
            file_stub = f"grade{w_grade}_{w_topic.replace(' ', '_')}_{w_kind}"

            dl_col1, dl_col2, dl_col3 = st.columns(3)
            with dl_col1:
                st.download_button("📥 .txt", data=content, file_name=f"{file_stub}.txt", mime="text/plain")
            with dl_col2:
                st.download_button(
                    "📥 Word (.docx)",
                    data=text_to_docx_bytes(export_text),
                    file_name=f"{file_stub}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            with dl_col3:
                st.download_button(
                    "📥 PDF", data=text_to_pdf_bytes(export_text), file_name=f"{file_stub}.pdf", mime="application/pdf"
                )

            with st.expander("📚 View Source Content"):
                st.write(topic_text)

# ==========================================================================
# TAB 3 — Coverage tracker
# ==========================================================================
with tab_coverage:
    st.caption("Track which grammar topics and stories you've already taught. Progress is saved automatically, per teacher.")

    coverage_key = f"coverage_state_{teacher_name}"
    if coverage_key not in st.session_state:
        st.session_state[coverage_key] = load_coverage_state(teacher_name)

    trackable_df = chunks_df[chunks_df["doc_type"].isin(["grammar_concept", "literature"])]
    coverage_grades = sorted(trackable_df["grade"].dropna().unique())

    total_items = len(trackable_df)
    total_done = sum(1 for cid in trackable_df["chunk_id"] if st.session_state[coverage_key].get(cid))
    if total_items:
        st.progress(total_done / total_items, text=f"Overall progress: {total_done}/{total_items} items covered")

    state_changed = False
    for g in coverage_grades:
        grade_items = trackable_df[trackable_df["grade"] == g]
        grade_done = sum(1 for cid in grade_items["chunk_id"] if st.session_state[coverage_key].get(cid))
        with st.expander(f"Grade {int(g)}  ·  {grade_done}/{len(grade_items)} covered"):
            for _, item in grade_items.iterrows():
                cid = item["chunk_id"]
                label = f"{'📖' if item['doc_type'] == 'literature' else '📘'} {item['title']}"
                current_value = st.session_state[coverage_key].get(cid, False)
                new_value = st.checkbox(label, value=current_value, key=f"cov_{teacher_name}_{cid}")
                if new_value != current_value:
                    st.session_state[coverage_key][cid] = new_value
                    state_changed = True

    if state_changed:
        save_coverage_state(teacher_name, st.session_state[coverage_key])

    st.divider()
    if st.button("🗑️ Reset coverage progress"):
        st.session_state[coverage_key] = {}
        save_coverage_state(teacher_name, {})
        st.rerun()

    st.divider()
    st.subheader("Export Coverage Report")
    report_lines = ["# Curriculum Coverage Report", "", f"Teacher: {teacher_name}", f"Overall progress: {total_done}/{total_items} items covered", ""]
    for g in coverage_grades:
        grade_items = trackable_df[trackable_df["grade"] == g]
        grade_done = sum(1 for cid in grade_items["chunk_id"] if st.session_state[coverage_key].get(cid))
        report_lines.append(f"## Grade {int(g)}  ({grade_done}/{len(grade_items)} covered)")
        for _, item in grade_items.iterrows():
            mark = "[x]" if st.session_state[coverage_key].get(item["chunk_id"]) else "[ ]"
            report_lines.append(f"{mark} {item['title']}")
        report_lines.append("")
    report_text = "\n".join(report_lines)

    rep_col1, rep_col2 = st.columns(2)
    with rep_col1:
        st.download_button(
            "📥 Word (.docx)",
            data=text_to_docx_bytes(report_text),
            file_name="curriculum_coverage_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    with rep_col2:
        st.download_button(
            "📥 PDF",
            data=text_to_pdf_bytes(report_text),
            file_name="curriculum_coverage_report.pdf",
            mime="application/pdf",
        )
